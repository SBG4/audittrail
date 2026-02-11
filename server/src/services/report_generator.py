"""Report data collection and document generation service.

Provides a shared data collection pipeline for all report formats (PDF, DOCX, HTML)
and rendering functions for each format.
"""

import asyncio
import io
import uuid
from functools import partial
from pathlib import Path
from datetime import datetime, timezone

from docx import Document
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.shared import Cm, Pt, RGBColor
from fastapi import HTTPException, status
from jinja2 import Environment, FileSystemLoader, select_autoescape
from sqlalchemy import select
from sqlalchemy.ext.asyncio import AsyncSession
from sqlalchemy.orm import selectinload
from weasyprint import HTML

from src.models.case import Case
from src.models.event import Event
from src.models.user import User

# Template setup
_TEMPLATES_DIR = Path(__file__).resolve().parent.parent / "templates"
_jinja_env = Environment(
    loader=FileSystemLoader(str(_TEMPLATES_DIR)),
    autoescape=select_autoescape(["html", "xml"]),
)

# Template name mapping
_TEMPLATE_MAP = {
    "timeline": "reports/timeline.html",
    "narrative": "reports/narrative.html",
}


async def collect_report_data(
    case_id: uuid.UUID,
    db: AsyncSession,
    current_user: User,
) -> dict:
    """Collect all data needed for report generation.

    Fetches the case with audit type, events with file batches,
    and formats everything into a template-ready dictionary.
    """
    # Fetch case with relationships
    result = await db.execute(
        select(Case)
        .where(Case.id == case_id)
        .options(
            selectinload(Case.audit_type),
            selectinload(Case.assigned_to),
            selectinload(Case.created_by),
        )
    )
    case = result.scalar_one_or_none()
    if case is None:
        raise HTTPException(
            status_code=status.HTTP_404_NOT_FOUND,
            detail="Case not found",
        )

    # Fetch events with file batches, sorted chronologically
    events_result = await db.execute(
        select(Event)
        .where(Event.case_id == case_id)
        .options(
            selectinload(Event.file_batches),
            selectinload(Event.created_by),
        )
        .order_by(
            Event.event_date.asc(),
            Event.event_time.asc().nulls_last(),
            Event.sort_order.asc(),
        )
    )
    events = events_result.scalars().all()

    # Build events data
    events_data = []
    total_file_count = 0
    for event in events:
        batches_data = []
        for batch in event.file_batches:
            batches_data.append({
                "label": batch.label,
                "file_count": batch.file_count,
                "description": batch.description or "",
                "file_types": batch.file_types or "",
            })

        event_file_count = event.file_count or 0
        total_file_count += event_file_count

        events_data.append({
            "event_date": event.event_date.isoformat() if event.event_date else "",
            "event_date_formatted": event.event_date.strftime("%Y-%m-%d") if event.event_date else "N/A",
            "event_time": event.event_time.strftime("%H:%M") if event.event_time else "",
            "event_time_formatted": event.event_time.strftime("%H:%M") if event.event_time else "N/A",
            "event_type": event.event_type or "note",
            "file_name": event.file_name or "",
            "file_count": event_file_count,
            "file_description": event.file_description or "",
            "file_type": event.file_type or "",
            "file_batches": batches_data,
            "has_batches": len(batches_data) > 0,
        })

    # Compute summary stats
    first_date = events[0].event_date.strftime("%Y-%m-%d") if events else "N/A"
    last_date = events[-1].event_date.strftime("%Y-%m-%d") if events else "N/A"

    # Count events by type
    findings_count = sum(1 for e in events_data if e["event_type"] == "finding")
    actions_count = sum(1 for e in events_data if e["event_type"] == "action")
    notes_count = sum(1 for e in events_data if e["event_type"] == "note")

    # Build metadata fields list
    metadata_fields = []
    if case.metadata_:
        for key, value in case.metadata_.items():
            metadata_fields.append({
                "label": key.replace("_", " ").title(),
                "value": str(value) if value else "N/A",
            })

    now = datetime.now(timezone.utc)

    return {
        "case": {
            "case_number": case.case_number,
            "title": case.title,
            "description": case.description or "",
            "status": case.status.title(),
            "created_at": case.created_at.strftime("%Y-%m-%d %H:%M") if case.created_at else "N/A",
            "updated_at": case.updated_at.strftime("%Y-%m-%d %H:%M") if case.updated_at else "N/A",
        },
        "audit_type": {
            "name": case.audit_type.name if case.audit_type else "Unknown",
            "slug": case.audit_type.slug if case.audit_type else "unknown",
        },
        "assigned_to": case.assigned_to.full_name if case.assigned_to else "Unassigned",
        "created_by": case.created_by.full_name if case.created_by else "Unknown",
        "metadata_fields": metadata_fields,
        "events": events_data,
        "stats": {
            "total_events": len(events_data),
            "total_file_count": total_file_count,
            "findings_count": findings_count,
            "actions_count": actions_count,
            "notes_count": notes_count,
            "first_date": first_date,
            "last_date": last_date,
        },
        "generated_at": now.strftime("%Y-%m-%d %H:%M UTC"),
        "generated_by": current_user.full_name,
    }


def _generate_pdf_sync(mode: str, data: dict) -> bytes:
    """Render HTML template and convert to PDF using WeasyPrint (sync).

    This is CPU-bound and must be called via run_in_executor.
    """
    template_name = _TEMPLATE_MAP.get(mode, _TEMPLATE_MAP["timeline"])
    template = _jinja_env.get_template(template_name)
    html_string = template.render(**data)
    html_doc = HTML(
        string=html_string,
        base_url=str(_TEMPLATES_DIR / "reports"),
    )
    return html_doc.write_pdf()


async def generate_pdf(mode: str, data: dict) -> bytes:
    """Generate a PDF report asynchronously.

    Runs WeasyPrint in a thread pool executor to avoid blocking the event loop.
    """
    loop = asyncio.get_event_loop()
    return await loop.run_in_executor(
        None, partial(_generate_pdf_sync, mode, data)
    )


# ---------------------------------------------------------------------------
# DOCX Generation
# ---------------------------------------------------------------------------

# Colors matching the PDF template
_HEADER_BG = RGBColor(44, 62, 80)  # #2c3e50
_HEADER_FG = RGBColor(255, 255, 255)
_FINDING_COLOR = RGBColor(231, 76, 60)  # #e74c3c
_ACTION_COLOR = RGBColor(52, 152, 219)  # #3498db
_NOTE_COLOR = RGBColor(149, 165, 166)  # #95a5a6
_LABEL_BG = RGBColor(245, 246, 250)  # #f5f6fa


def _set_cell_shading(cell, color_hex: str) -> None:
    """Set background shading on a table cell."""
    shading = cell._element.get_or_add_tcPr()
    shading_elm = shading.makeelement(
        qn("w:shd"),
        {
            qn("w:val"): "clear",
            qn("w:color"): "auto",
            qn("w:fill"): color_hex,
        },
    )
    shading.append(shading_elm)


def _style_header_row(table, num_cols: int) -> None:
    """Style the first row of a table as a header with dark background and white text."""
    for i in range(num_cols):
        cell = table.rows[0].cells[i]
        _set_cell_shading(cell, "2C3E50")
        para = cell.paragraphs[0]
        if para.runs:
            run = para.runs[0]
        else:
            run = para.add_run(para.text)
            para.text = ""
        run.bold = True
        run.font.color.rgb = _HEADER_FG
        run.font.size = Pt(9)


def _add_metadata_table(doc: Document, data: dict) -> None:
    """Add a two-column metadata table to the document."""
    rows_data = [
        ("Case Number", f"#{data['case']['case_number']}"),
        ("Title", data["case"]["title"]),
        ("Audit Type", data["audit_type"]["name"]),
        ("Status", data["case"]["status"]),
        ("Assigned To", data["assigned_to"]),
        ("Created By", data["created_by"]),
        ("Created At", data["case"]["created_at"]),
        ("Updated At", data["case"]["updated_at"]),
    ]
    # Add custom metadata fields
    for field in data.get("metadata_fields", []):
        rows_data.append((field["label"], field["value"]))

    table = doc.add_table(rows=len(rows_data), cols=2, style="Table Grid")
    table.alignment = WD_TABLE_ALIGNMENT.CENTER

    for i, (label, value) in enumerate(rows_data):
        # Label cell
        label_cell = table.rows[i].cells[0]
        label_cell.paragraphs[0].text = label
        label_para = label_cell.paragraphs[0]
        if label_para.runs:
            label_para.runs[0].bold = True
            label_para.runs[0].font.size = Pt(9)
        else:
            run = label_para.add_run(label)
            label_para.clear()
            label_para.add_run(label).bold = True
            label_para.runs[0].font.size = Pt(9)
        _set_cell_shading(label_cell, "F5F6FA")

        # Value cell
        value_cell = table.rows[i].cells[1]
        value_cell.paragraphs[0].text = str(value)
        if value_cell.paragraphs[0].runs:
            value_cell.paragraphs[0].runs[0].font.size = Pt(9)

    # Set column widths
    for row in table.rows:
        row.cells[0].width = Cm(5)
        row.cells[1].width = Cm(12)


def _add_events_table(doc: Document, events: list[dict]) -> None:
    """Add a chronological events table to the document."""
    headers = ["Date", "Time", "Type", "File Name", "Count", "Description"]
    table = doc.add_table(rows=1, cols=len(headers), style="Table Grid")
    table.alignment = WD_TABLE_ALIGNMENT.CENTER

    # Set header row
    for i, header in enumerate(headers):
        cell = table.rows[0].cells[i]
        cell.paragraphs[0].text = header
    _style_header_row(table, len(headers))

    # Add event rows
    for event in events:
        row = table.add_row()
        row.cells[0].paragraphs[0].text = event["event_date_formatted"]
        row.cells[1].paragraphs[0].text = event["event_time_formatted"]
        row.cells[2].paragraphs[0].text = event["event_type"].upper()
        row.cells[3].paragraphs[0].text = event["file_name"] or "\u2014"
        row.cells[4].paragraphs[0].text = str(event["file_count"]) if event["file_count"] else "\u2014"
        row.cells[5].paragraphs[0].text = event["file_description"] or "\u2014"

        # Style type cell with color
        type_run = row.cells[2].paragraphs[0].runs[0] if row.cells[2].paragraphs[0].runs else row.cells[2].paragraphs[0].add_run(event["event_type"].upper())
        if event["event_type"] == "finding":
            type_run.font.color.rgb = _FINDING_COLOR
        elif event["event_type"] == "action":
            type_run.font.color.rgb = _ACTION_COLOR
        else:
            type_run.font.color.rgb = _NOTE_COLOR
        type_run.bold = True
        type_run.font.size = Pt(8)

        # Set font size for all cells
        for cell in row.cells:
            for para in cell.paragraphs:
                for run in para.runs:
                    run.font.size = Pt(9)

        # Add file batch sub-rows if present
        if event["has_batches"]:
            batch_row = table.add_row()
            batch_row.cells[0].merge(batch_row.cells[5])
            merged_cell = batch_row.cells[0]
            merged_cell.paragraphs[0].text = ""
            para = merged_cell.paragraphs[0]
            run = para.add_run("File Batches: ")
            run.bold = True
            run.font.size = Pt(8)
            run.font.color.rgb = RGBColor(85, 85, 85)
            for batch in event["file_batches"]:
                batch_text = f"\n    {batch['label']} \u2014 {batch['file_count']} file(s)"
                if batch["file_types"]:
                    batch_text += f" | Types: {batch['file_types']}"
                if batch["description"]:
                    batch_text += f" | {batch['description']}"
                run = para.add_run(batch_text)
                run.font.size = Pt(8)
                run.font.color.rgb = RGBColor(85, 85, 85)


def _add_numbered_items(
    doc: Document,
    events: list[dict],
    event_type: str,
    color: RGBColor,
) -> None:
    """Add numbered items for a specific event type (findings, actions, notes)."""
    filtered = [e for e in events if e["event_type"] == event_type]

    if not filtered:
        para = doc.add_paragraph()
        run = para.add_run(f"No {event_type}s recorded.")
        run.italic = True
        run.font.color.rgb = RGBColor(189, 195, 199)
        return

    for idx, event in enumerate(filtered, 1):
        # Header line
        para = doc.add_paragraph()
        header_text = f"{event_type.title()} {idx} \u2014 {event['event_date_formatted']}"
        if event["event_time_formatted"] != "N/A":
            header_text += f" at {event['event_time_formatted']}"
        run = para.add_run(header_text)
        run.bold = True
        run.font.size = Pt(10)
        run.font.color.rgb = color

        # File details
        if event["file_name"]:
            detail_para = doc.add_paragraph()
            run = detail_para.add_run(f"File: {event['file_name']}")
            run.font.size = Pt(9)
            if event["file_count"]:
                run = detail_para.add_run(f" ({event['file_count']} file(s))")
                run.font.size = Pt(9)
            if event["file_type"]:
                run = detail_para.add_run(f" | Type: {event['file_type']}")
                run.font.size = Pt(9)

        # Description
        if event["file_description"]:
            desc_para = doc.add_paragraph()
            run = desc_para.add_run(event["file_description"])
            run.font.size = Pt(9)

        # File batches
        if event["has_batches"]:
            batch_para = doc.add_paragraph()
            run = batch_para.add_run("Associated File Batches:")
            run.bold = True
            run.font.size = Pt(8)
            for batch in event["file_batches"]:
                batch_text = f"\n    {batch['label']} \u2014 {batch['file_count']} file(s)"
                if batch["file_types"]:
                    batch_text += f" | Types: {batch['file_types']}"
                if batch["description"]:
                    batch_text += f" | {batch['description']}"
                run = batch_para.add_run(batch_text)
                run.font.size = Pt(8)

        # Add spacing
        doc.add_paragraph()


def _generate_docx_timeline(data: dict) -> bytes:
    """Generate a DOCX report in quick timeline mode (sync)."""
    doc = Document()

    # Title
    title = doc.add_heading("Audit Case Report", level=0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER

    subtitle = doc.add_heading("Quick Timeline", level=1)
    subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER

    # Case info line
    para = doc.add_paragraph()
    para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = para.add_run(f"Case #{data['case']['case_number']}: {data['case']['title']}")
    run.font.size = Pt(12)
    run.font.color.rgb = RGBColor(52, 73, 94)

    para = doc.add_paragraph()
    para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = para.add_run(f"Generated {data['generated_at']}")
    run.font.size = Pt(9)
    run.font.color.rgb = RGBColor(149, 165, 166)

    doc.add_page_break()

    # Case Summary
    doc.add_heading("Case Summary", level=2)
    _add_metadata_table(doc, data)

    # Description
    if data["case"]["description"]:
        doc.add_heading("Description", level=3)
        para = doc.add_paragraph()
        run = para.add_run(data["case"]["description"])
        run.font.size = Pt(10)

    # Timeline Events
    doc.add_heading("Timeline Events", level=2)

    if data["events"]:
        _add_events_table(doc, data["events"])
    else:
        para = doc.add_paragraph()
        run = para.add_run("No events recorded for this case.")
        run.italic = True
        run.font.color.rgb = RGBColor(189, 195, 199)

    # Footer
    doc.add_paragraph()
    footer_para = doc.add_paragraph()
    footer_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = footer_para.add_run(
        f"Generated on {data['generated_at']} by {data['generated_by']} \u2014 AuditTrail"
    )
    run.font.size = Pt(8)
    run.font.color.rgb = RGBColor(153, 153, 153)

    # Save to bytes
    buffer = io.BytesIO()
    doc.save(buffer)
    buffer.seek(0)
    return buffer.getvalue()


def _generate_docx_narrative(data: dict) -> bytes:
    """Generate a DOCX report in detailed narrative mode (sync)."""
    doc = Document()

    # Title
    title = doc.add_heading("Audit Case Report", level=0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER

    subtitle = doc.add_heading("Detailed Narrative", level=1)
    subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER

    # Case info line
    para = doc.add_paragraph()
    para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = para.add_run(f"Case #{data['case']['case_number']}: {data['case']['title']}")
    run.font.size = Pt(12)
    run.font.color.rgb = RGBColor(52, 73, 94)

    para = doc.add_paragraph()
    para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = para.add_run(f"Generated {data['generated_at']}")
    run.font.size = Pt(9)
    run.font.color.rgb = RGBColor(149, 165, 166)

    doc.add_page_break()

    # Executive Summary
    doc.add_heading("Executive Summary", level=2)

    stats = data["stats"]
    case = data["case"]

    para = doc.add_paragraph()
    run = para.add_run(
        f"This report covers audit case #{case['case_number']} "
        f"({case['title']}), a {data['audit_type']['name']} investigation. "
        f"The case currently has a status of {case['status']} "
        f"and is assigned to {data['assigned_to']}."
    )
    run.font.size = Pt(10)

    if stats["total_events"] > 0:
        para = doc.add_paragraph()
        run = para.add_run(
            f"The timeline spans {stats['total_events']} event(s) recorded "
            f"between {stats['first_date']} and {stats['last_date']}, "
            f"involving a total of {stats['total_file_count']} file(s)."
        )
        run.font.size = Pt(10)
    else:
        para = doc.add_paragraph()
        run = para.add_run("No events have been recorded for this case yet.")
        run.font.size = Pt(10)

    # Key metrics table
    metrics_table = doc.add_table(rows=2, cols=5, style="Table Grid")
    metrics_table.alignment = WD_TABLE_ALIGNMENT.CENTER
    metric_labels = ["Total Events", "Findings", "Actions", "Notes", "Total Files"]
    metric_values = [
        str(stats["total_events"]),
        str(stats["findings_count"]),
        str(stats["actions_count"]),
        str(stats["notes_count"]),
        str(stats["total_file_count"]),
    ]
    for i in range(5):
        # Value row
        cell = metrics_table.rows[0].cells[i]
        cell.paragraphs[0].text = metric_values[i]
        cell.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
        if cell.paragraphs[0].runs:
            cell.paragraphs[0].runs[0].bold = True
            cell.paragraphs[0].runs[0].font.size = Pt(14)
        # Label row
        cell = metrics_table.rows[1].cells[i]
        cell.paragraphs[0].text = metric_labels[i]
        cell.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
        if cell.paragraphs[0].runs:
            cell.paragraphs[0].runs[0].font.size = Pt(8)
            cell.paragraphs[0].runs[0].font.color.rgb = RGBColor(127, 140, 141)

    # Case Details
    doc.add_heading("Case Details", level=2)
    _add_metadata_table(doc, data)

    if data["case"]["description"]:
        doc.add_heading("Description", level=3)
        para = doc.add_paragraph()
        run = para.add_run(data["case"]["description"])
        run.font.size = Pt(10)

    # Findings
    doc.add_page_break()
    doc.add_heading("Findings", level=2)
    _add_numbered_items(doc, data["events"], "finding", _FINDING_COLOR)

    # Actions Taken
    doc.add_heading("Actions Taken", level=2)
    _add_numbered_items(doc, data["events"], "action", _ACTION_COLOR)

    # Supporting Notes
    doc.add_heading("Supporting Notes", level=2)
    _add_numbered_items(doc, data["events"], "note", _NOTE_COLOR)

    # Complete Timeline
    doc.add_page_break()
    doc.add_heading("Complete Timeline", level=2)

    if data["events"]:
        _add_events_table(doc, data["events"])
    else:
        para = doc.add_paragraph()
        run = para.add_run("No events recorded for this case.")
        run.italic = True
        run.font.color.rgb = RGBColor(189, 195, 199)

    # Conclusions
    doc.add_heading("Conclusions", level=2)
    para = doc.add_paragraph()
    run = para.add_run("[Add conclusions here]")
    run.italic = True
    run.font.color.rgb = RGBColor(189, 195, 199)

    # Recommendations
    doc.add_heading("Recommendations", level=2)
    para = doc.add_paragraph()
    run = para.add_run("[Add recommendations here]")
    run.italic = True
    run.font.color.rgb = RGBColor(189, 195, 199)

    # Footer
    doc.add_paragraph()
    footer_para = doc.add_paragraph()
    footer_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = footer_para.add_run(
        f"Generated on {data['generated_at']} by {data['generated_by']} \u2014 AuditTrail"
    )
    run.font.size = Pt(8)
    run.font.color.rgb = RGBColor(153, 153, 153)

    # Save to bytes
    buffer = io.BytesIO()
    doc.save(buffer)
    buffer.seek(0)
    return buffer.getvalue()


async def generate_docx(mode: str, data: dict) -> bytes:
    """Generate a DOCX report asynchronously.

    Runs python-docx generation in a thread pool executor.
    """
    if mode == "narrative":
        func = _generate_docx_narrative
    else:
        func = _generate_docx_timeline

    loop = asyncio.get_event_loop()
    return await loop.run_in_executor(None, partial(func, data))
